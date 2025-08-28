# Plankton code
import pandas as pd
import geopandas as gpd
import matplotlib.pyplot as plt
import seaborn as sns
from shapely.geometry import LineString
from docx import Document
from docx.shared import Inches
import shutil
from fastai.vision.all import *
import os
from PIL import Image
from PIL.ExifTags import GPSTAGS
from matplotlib.ticker import FuncFormatter, MultipleLocator
import matplotlib.image as mpimg
import polars as pl

# Custom modules
from src.report_visualizations import *

def get_pred_labels(TRAIN_DATA_PATH, MODEL_FILENAME):
    # Get label names
    # Quite convoluted, but fastest way to retrieve labels in a dynamic way
    block = DataBlock(
        blocks=(ImageBlock, CategoryBlock),
        splitter=RandomSplitter(),
        get_items=get_image_files,
        get_y=parent_label,
        item_tfms=Resize(300, ResizeMethod.Pad, pad_mode='zeros'),
        batch_tfms=[*aug_transforms(
            mult=1.0,
            do_flip=True,
            flip_vert=True,
            max_rotate=0.2,
            min_zoom=1.0,
            max_zoom=1.1,
            max_lighting=0.3,
            max_warp=0.1,
            p_affine=0.5,
            p_lighting=0.5,
            pad_mode='zeros'),
            Normalize.from_stats(*imagenet_stats)]
    )

    dls = block.dataloaders(TRAIN_DATA_PATH, bs=600)
    learn = vision_learner(dls, resnet50, metrics=error_rate)
    learn.load(MODEL_FILENAME, weights_only=False)

    # Actual label names extracted from the model
    global pred_labels # Make global to use in other functions
    pred_labels = learn.dls.vocab # Extract label names as list
    print(f"[INFO] Prediction labels available in {MODEL_FILENAME}:\n{pred_labels}")

    return pred_labels

def get_geographic_data(image_path):
    def convert_to_degrees(value):
        """Convert GPS coordinates to decimal degrees."""
        degrees, minutes, seconds = value
        return degrees + (minutes / 60.0) + (seconds / 3600.0)

    # Extract latitude-longitude from the EXIF metadata
    with Image.open(image_path) as image:
        exif_data = image.getexif()
        if exif_data:
            ifd = exif_data.get_ifd(0x8825) # Code for GPSInfo, see: www.media.mit.edu/pia/Research/deepview/exif.html
            
            # TODO" Check for (0,0) coordinates
            if not ifd:
                print(f"[WARNING] '{image_path}' has no GPS information.")
                return None

            gps_info = {}
            for key, val in ifd.items():
                gps_info[GPSTAGS.get(key, key)] = val

            # Extract and convert latitude and longitude
            latitude = convert_to_degrees(gps_info['GPSLatitude'])
            if gps_info['GPSLatitudeRef'] != 'N':
                latitude = -latitude

            longitude = convert_to_degrees(gps_info['GPSLongitude'])
            if gps_info['GPSLongitudeRef'] != 'E':
                longitude = -longitude

            return latitude, longitude

# General data description
def compute_class_statistics(df, total_images, DENSITY_CONSTANT):
    # Get unique class IDs
    class_ids = df['pred_id'].unique()

    # Compute statistics for each class
    for class_id in class_ids:
        class_df = df.filter(pl.col("pred_id") == class_id)

        total_class_images = class_df.height
        if total_class_images == 0:
            continue # Skip classes with no images
        percentage_total = (total_class_images / total_images) * 100

        # Collect per-class statistics
        stats_dict = ({
            'ID': class_id,
            'Class': class_df["pred_label"][0],  # Return as single string
            '# of Images': total_class_images,
            '% of Images': f"{percentage_total:.2f}",
            'Min Density': f"{class_df['density'].min():.2f}",
            'Mean Density': f"{class_df['density'].mean():.2f}",
            'Max Density': f"{class_df['density'].max():.2f}",
            'Min Confidence': f"{class_df['pred_conf'].min() * 100:.2f}", # Express as %
            'Mean Confidence': f"{class_df['pred_conf'].mean() * 100:.2f}",
            'Max Confidence': f"{class_df['pred_conf'].max() * 100:.2f}"
        })

    return stats_dict

# Per-class data summaries
def preprocess_data(df, class_id, DENSITY_CONSTANT):
    # Filter the DataFrame for the specified class
    class_df = df.filter(pl.col("pred_id") == class_id).sort(by='datetime')

    total_class_images = class_df.height
    if total_class_images == 0:
        print(f"[INFO] No images were classified as {class_df['pred_label'].first()}.")
        return

    # General statistics
    total_images = df.height
    percentage_total = (total_class_images / total_images) * 100

    # Group by 10-minute intervals and aggregate
    grouped_df = (
        class_df
        .with_columns(pl.col("datetime").str.to_datetime())  # Convert 'datetime' to datetime type
        .group_by_dynamic("datetime", every="10m")
        .agg([
            pl.col("id").first(),
            pl.col("date").first(),
            pl.col("time").first(),
            pl.col("pred_conf").mean().alias("pred_conf"),
            pl.col("pred_id").first(),
            pl.col("pred_label").first(),
            pl.col("density").mean()
            # (pl.len() / DENSITY_CONSTANT).alias("density") # Density in N/L
        ])
    )

    # Create a simple DataFrame to pretty-print the output statistics
    # Note the two different DataFrames used, with the density statistics being grouped per 10 minute bins
    data = pd.DataFrame({
        "Metric": [
            "Class ID",
            'Class',
            "Number of predictions",
            "Percentage of total images",
            "Min confidence",
            "Max confidence",
            "Average confidence",
            "Min density",
            "Max density",
            "Average density"
        ],
        "Value": [
            class_id,
            grouped_df['pred_label'].first(),
            f"{total_class_images:,}",
            f"{percentage_total:.2f} %",
            f"{class_df['pred_conf'].min() * 100:.2f} %",
            f"{class_df['pred_conf'].max() * 100:.2f} %",
            f"{class_df['pred_conf'].mean() * 100:.2f} %",
            f"{grouped_df['density'].min():.2f} n/L",
            f"{grouped_df['density'].max():.2f} n/L",
            f"{grouped_df['density'].mean():.2f} n/L"
        ]
    })

    return class_df, grouped_df, data

# Compute entire cruise-path
def create_cruise_path(lazy_df, CRUISE_NAME):
    print("[INFO] Creating GeoDataFrame of cruise path")

    # Group by datetime and get first lat and lon for each group
    measurement_locations = (
        lazy_df
        .group_by("datetime")
        .agg([
            pl.col("lat").first().alias("lat"),
            pl.col("lon").first().alias("lon")
        ])
    ).collect().to_pandas().set_index('datetime')

    # Create a GeoDataFrame from the adjusted data
    geometry = gpd.points_from_xy(measurement_locations['lon'], measurement_locations['lat'], crs=4326)
    gdf = gpd.GeoDataFrame(measurement_locations, geometry=geometry)

    # Sort the GeoDataFrame by the index (datetime) to ensure the points are in chronological order
    gdf.sort_index(inplace=True)
    gdf.index = pd.to_datetime(gdf.index)

    # Save measurement locations as GeoPackage
    output_path = Path(f"data/{CRUISE_NAME}_results/{CRUISE_NAME}_measurement_locations.gpkg")
    gdf.to_file(output_path, driver='GPKG')

    # Create a linestring from the sorted points
    cruise_path = LineString(gdf['geometry'].tolist())

    # Calculate time differences between consecutive points
    datetimes = gdf.index.to_series()
    time_diffs = (datetimes - datetimes.shift(1)).dt.total_seconds() / 60
    time_diffs.iloc[0] = 0  # First point has no previous point to compare to

    # Identify gaps > 10 minutes
    gap_locations = time_diffs > 10

    # Calculate segment IDs by cumulative sum of gap locations
    segment_ids = gap_locations.cumsum()

    # Add segment IDs to the GeoDataFrame
    gdf['segment_id'] = segment_ids

    # Group by segment_id and create linestrings for groups with at least 2 points
    segments = []
    for segment_id, group in gdf.groupby('segment_id'):
        if len(group) >= 2:
            segments.append(LineString(group['geometry'].tolist()))

    # Create a GeoDataFrame for the cruise path segments
    cruise_path_gdf = gpd.GeoDataFrame(geometry=segments, crs="EPSG:4326")
    
    # Reproject to EPSG:3035
    cruise_path_gdf = cruise_path_gdf.to_crs(epsg=3035) # CRS with meters

    cruise_path_output_path = Path(f"data/{CRUISE_NAME}_results/{CRUISE_NAME}_cruise_path.gpkg")
    cruise_path_gdf.to_file(cruise_path_output_path, driver='GPKG')

    # Other statics used for pretty-print and general information
    cruise_path_gdf['length_km'] = cruise_path_gdf['geometry'].to_crs(epsg=3035).length / 1000
    total_length_km = cruise_path_gdf['length_km'].sum()
    print(f"[INFO] Total length of all linestring segments: {total_length_km:.2f} km")

    # Calculate total hours of footage
    number_of_samples = len(gdf)
    total_hours = (number_of_samples * 10) / 60  # 10 minutes per sample, divided by 60 to get hours
    print(f"[INFO] Total hours of footage: {total_hours:.2f}")

    # Get min and max dates
    min_date = gdf.index.min()
    max_date = gdf.index.max()
    print(f"[INFO] Minimum date: {min_date}")
    print(f"[INFO] Maximum date: {max_date}")

    cruise_path_gdf = cruise_path_gdf.to_crs(epsg=4258) # Use CRS with degrees for plotting purposes

    return cruise_path_gdf, total_length_km, total_hours, min_date, max_date

# Automated report
# @profile
def create_word_document(results_dir, OSPAR, CRUISE_NAME, DENSITY_CONSTANT, TRAIN_DATASET, MODEL_FILENAME):
    print(f"[INFO] Reading DataFrames in folder: {results_dir}")

    # To reduce memory load from ~80GB CSV files, we use Polars + LazyFrames
    # First create glob pattern to find available .csv files
    csv_files = list(Path(results_dir).glob("*.csv"))

    if not csv_files:
        raise FileNotFoundError(f"[ERROR] No CSV files found in {results_dir}")
    lazy_df = pl.concat([pl.scan_csv(str(file)) for file in csv_files])
    
    # Check if 'label' column exists and rename only if necessary
    if "label" in lazy_df.collect_schema().keys():
        lazy_df = lazy_df.rename({"label": "pred_id"})
        print("[INFO] Renamed 'label' column to 'pred_id' for backwards compatibility")
    else:
        print("[INFO] No 'label' column found, continuing with 'pred_id'")

    # Then load in essential information to dynamically loop over the data later on
    total_rows = lazy_df.select(pl.len()).collect().item()
    total_classes = lazy_df.select(pl.col("pred_id").unique()).collect().to_series().to_list()
    print(f"[INFO] Read DataFrame. Started processing {total_rows:,} rows.")

    # Get prediction labels for post-processing
    # pred_labels is a global variable (for now)
    pred_labels = get_pred_labels(TRAIN_DATASET, MODEL_FILENAME)

    # Initialize data structures
    class_stats = [] # To store statistics in general table
    class_data = {} # To store per-class data for document generation

    # Create temporary directory for figures
    temp_dir = 'data/temp'
    os.makedirs(temp_dir, exist_ok=True)

    # Get cruise path information as geodata
    cruise_path, total_length_km, total_hours, start_date, end_date = create_cruise_path(lazy_df, CRUISE_NAME) # GeoDataFrame
    minx, miny, maxx, maxy = cruise_path.total_bounds # Used for general text description

    def longitude_direction(lon):
        # Determine the direction for longitude
        return "W" if lon < 0 else "E"

    # Iterate over each class to read in less volume, compared ot the entire dataset at once
    # Can still be reasonably high (>10,000,000 rows) with some classes (e.g., detritus)
    for class_id in total_classes:
        print(f"{'-' * 50}")
        print(f"[INFO] Processing class {class_id}")

        # Get and clean subset for current class
        subset_df = lazy_df.filter(pl.col("pred_id") == class_id).collect()
        print(f"[INFO] Processing {subset_df.height:,} rows for class {class_id}")

        # 1. Compute and store statistics (number of predictions, density, model confidence)
        stats_dict = compute_class_statistics(subset_df, total_rows, DENSITY_CONSTANT)
        class_stats.append(stats_dict) # Single row per class

        # 2. Generate figure on model confidence compared to top-5 related classes
        confidence_fig = plot_confidence(subset_df)

        # 3. Generate figure on density statistics over time
        density_fig = plot_density_graph(subset_df, class_id, pred_labels)

        # 4. Generate map of density on a spatial scale
        map_fig = plot_class_density_map(subset_df, class_id, pred_labels, cruise_path, CRUISE_NAME)

        # 5. Generate figure of randomly selected images to illustrate predicted targets
        img_fig = plot_random_images(subset_df, num_images=80)

        if any(fig is None for fig in [density_fig, map_fig, confidence_fig]):
            print(f"[WARNING] Skipping class {class_id} due to missing figures")
            continue

        # Save figures and store paths
        figure_paths = {}
        for fig_type, fig in [
            ('confidence', confidence_fig),
            ('density', density_fig),
            ('map', map_fig),
            ('img', img_fig)
        ]:
            path = os.path.join(temp_dir, f'{fig_type}_fig_{class_id}.png')
            fig.savefig(path)
            plt.close(fig)
            figure_paths[fig_type] = path

        # Store all data needed for document generation
        class_data[class_id] = {
            'pred_label': subset_df['pred_label'].first(),
            'figure_paths': figure_paths
        }

        # TODO: Remove
        if class_id > 3:
            break

    # Create Word document
    document = Document()
    document.add_heading(f'Automated report for {CRUISE_NAME} survey', 0)

    document.add_heading('Introduction', level=1)
    document.add_paragraph(
        "This automated report provides a detailed overview of Plankton Imager data collected during the " +
        f"{CRUISE_NAME} cruise between {str(start_date)} to {str(end_date)}, with a total of {total_hours:.2f} hours of footage recorded. In total, {total_rows:,} images were collected over a transect of {total_length_km:.2f} kilometers in the region of " +
        f"from approximately {abs(minx):.2f}°{longitude_direction(minx)} to {abs(maxx):.2f}°{longitude_direction(maxx)} longitude and " +
        f"{miny:.2f}°N to {maxy:.2f}°N latitude.\n\n" 

        "This automated report provides an overview of the data processed using the ResNet50 model developed by van Walraven et al. (in prep), see method section. "+
        "The report contains information on the number of images per class, density statistics (N/L), and the model confidence in predicting the class. " +
        "Moreover, figures are created on the 10-minute bins of detected objects visualized in density and spatio-temporal plots. " +
        f"A first attempt at deriving an index of patchiness is provided as well, through dividing the number of images by {DENSITY_CONSTANT}, corresponding to the volume of water (in L) flowing through " +
        "the Plankton Imager in 10 minutes."
    )

    # Plot map of the cruise-path
    document.add_heading('Cruise overview', level=2)
    cruise_fig_path = os.path.join(temp_dir, f'cruise_path.png')
    cruise_fig = plot_cruise_path(cruise_path, cruise_fig_path, CRUISE_NAME)
    document.add_picture(cruise_fig_path, width=Inches(6))
    document.add_paragraph(f"Map showing the path taken during the {CRUISE_NAME} survey.")

    # Plot graph of data availability during the cruise
    data_availability_path = create_data_availability_plot(lazy_df, 'datetime', start_date, end_date, temp_dir)
    document.add_picture(data_availability_path, width=Inches(6))
    document.add_paragraph(f"Data availability during the {CRUISE_NAME} survey.")

    document.add_heading('Methodology', level=1)
    document.add_paragraph(
        "The analysis was conducted using the ResNet50 developed in Van Walraven et al. (in prep) which predicts 49 different plankton and non-plankton classes. Detailed information " + 
        "on the code, weights, and datasets can be found at: https://github.com/geoJoost/plankton_imager_classifier. " +
        "The density calculations are created by dividing the estimated volume of water that passes through the flow cell 34 L/min, for a total of 340 L per 10 minutes. " + 
        "Therefore, the number of predictions can be binned per 10 minutes and divided to get the density per L, which we report as N/L. "
        "Afterwards, we process this dataset to compute descriptive statistics on the entire dataset, and class-specific derived information. Density graphs and maps are created to illustrate the "+
        "distribution of each predicted class over time and geographic locations.\n\n"

        "Moreover, to provide an estimate of the model's confidence in predicting specific classes, we generate descriptive statistics and a violin plot illustrating the confidence distribution. " +
        "This visualization helps in understanding the variability and uncertainty in the model's predictions." +
        "We focus on the top five classes most related to the actual target by computing the mean confidence values for all non-target classes and selecting " +
        "the five with the highest mean confidence. These classes are considered the most related to the target and represent areas where the model might struggle to differentiate between the actual target and similar classes. " +
        "For example, the model might show lower certainty in predictions for the class 'Crustacea Amphipoda', particularly when differentiating it from other closely related Crustacea targets. " +
        "This adds a layer of nuance to the model, and shows potential areas for improvements."
    )

    document.add_heading('General statistics', level=1)
    document.add_paragraph(
        f"During the {CRUISE_NAME} survey, a total of {total_rows:,} images were collected along the transect. "+
        "The table below shows the amount of images predicted for each class and as percentage of the total images collected. "+
        "Also shown is the minimum, mean, and maximum density (n/L) for each class."
    )

    # Add a table for density statistics
    document.add_paragraph("General statistics on the predictions made using the model, such as the number of images predicted for each class, density (N/L), and the confidence values.")
    
    # Aggregate statistics into single table
    stats_df = pd.DataFrame(class_stats).sort_values(by='Class', ascending=True) # 49 rows
    stats_df['Class'] = stats_df['Class'].str.replace('_', ' ').str.replace('-', ' ', regex=False) # Use this to make the column smaller

    # Create a table with 5 columns
    table = document.add_table(rows=1, cols=5)
    table.style = 'Light Shading Accent 1'

    # Set the width of each column
    column_widths = [Inches(0.5), Inches(2.5), Inches(1.1), Inches(1.1), Inches(1.1)]
    for i, width in enumerate(column_widths):
        table.columns[i].width = width

    # Add header row
    headers = ['ID', 'Class', '# of Images', 'Density', 'Confidence']
    for i, header in enumerate(headers):
        table.cell(0, i).text = header

    # Add data rows
    for _, row in stats_df.iterrows():
        # Add a new row for each class
        cells = table.add_row().cells
        cells[0].text = str(row['ID'])
        cells[1].text = row['Class']
        cells[2].text = f"{row['# of Images']:,}\n{row['% of Images']}%"

        # Add Density statistics as multiple lines in the Density column
        density_text = f"Min: {row['Min Density']}\nMean: {row['Mean Density']}\nMax: {row['Max Density']}"
        cells[3].text = density_text

        # Add Confidence statistics as multiple lines in the Confidence column
        confidence_text = f"Min: {row['Min Confidence']}\nMean: {row['Mean Confidence']}\nMax: {row['Max Confidence']}"
        cells[4].text = confidence_text
    
    # Plot bar-chart of the image counts
    count_path = os.path.join(temp_dir, f'count_graph.png')
    plot_image_count(stats_df, count_path)
    document.add_picture(count_path, width=Inches(6))
    document.add_paragraph(f"Graph showing the image counts of classes predicted during the {CRUISE_NAME} survey.")

    # Add page break before class-specific sections
    document.add_page_break()

    # Add per-class sections
    for class_id in stats_df['ID']:
        class_info = class_data[class_id]
        pred_label = class_info['pred_label']
        report_df = stats_df[stats_df['ID'] == class_id] # Use the generalized table to print statistics again
        figure_paths = class_info['figure_paths']

        document.add_heading(f'Class: {pred_label}', level=1)

        # Add statistics table
        document.add_heading('Statistics', level=2)
        table = document.add_table(rows=len(report_df.columns) + 1, cols=2)
        table.style = 'Light Shading Accent 1'

        # Add header row
        table.cell(0, 0).text = 'Metric'
        table.cell(0, 1).text = 'Value'

        # Map each DataFrame column to its display name and unit
        metrics_map = [
            ('ID', 'Class ID', ''),
            ('Class', 'Class', ''),
            ('# of Images', 'Number of predictions', ''),
            ('% of Images', 'Percentage of total images', '%'),
            ('Min Confidence', 'Min confidence', '%'),
            ('Max Confidence', 'Max confidence', '%'),
            ('Mean Confidence', 'Average confidence', '%'),
            ('Min Density', 'Min density', 'n/L'),
            ('Max Density', 'Max density', 'n/L'),
            ('Mean Density', 'Average density', 'n/L')
        ]

        # Populate the table
        for i, (df_col, display_name, unit) in enumerate(metrics_map, start=1):
            # Get the raw value
            value = report_df.iloc[0][df_col]

            # Format the value with its unit
            if unit == '%':
                # For percentages, we assume the value is already in percentage form (0-100)
                formatted_value = f"{value} %"
            elif unit == 'n/L':
                # For densities, format with 2 decimal places
                formatted_value = f"{value} {unit}"
            elif df_col == '# of Images':
                # Special case for counts
                formatted_value = f"{int(value):,}"
            else:
                # For all other values, just add the unit if there is one
                formatted_value = f"{value}{' ' + unit if unit else ''}"

            # Add to table
            table.cell(i, 0).text = display_name
            table.cell(i, 1).text = formatted_value

        # Add figures to the document
        document.add_heading('Model confidence', level=2)
        document.add_picture(figure_paths['confidence'], width=Inches(6))
        document.add_paragraph(f"Figure showcasing the confidence distribution of the top-5 most related classes, which are determined by computing the mean confidence value of all non-target classes and selecting the five with the highest means, compared to {pred_label}.")

        document.add_heading('Temporal density', level=2)
        document.add_picture(figure_paths['density'], width=Inches(6))
        document.add_paragraph(f"Figure showcasing the daily density estimates for {pred_label} per 10 minutes.")

        document.add_heading('Spatio-temporal density', level=2)
        document.add_picture(figure_paths['map'], width=Inches(6))
        document.add_paragraph(f"Map showing mean density estimates for {pred_label} along the cruise transect. OSPAR eutrophication areas are plotted in the background with different colours.")
        
        document.add_heading('Classification examples', level=2)
        document.add_picture(figure_paths['img'], width=Inches(6))
        document.add_paragraph(f"Figure showing several images classified as {pred_label} with model confidence.")

        # Close the figures to free memory
        plt.close(confidence_fig)
        plt.close(density_fig)
        plt.close(map_fig)
        plt.close(img_fig)

        document.add_page_break() # Ensure a new page starts for each class

    # Save the report
    document_path = f'reports/report_{CRUISE_NAME}.docx'
    os.makedirs(os.path.dirname(document_path), exist_ok=True)
    document.save(document_path)

    # Delete the temporary directory
    shutil.rmtree(temp_dir)
    print(f"[INFO] Word document created at: {document_path}")

    return document_path

# if __name__ == "__main__":
#     # Hard-coded variables
#     MODEL_FILENAME = Path('Plankton_imager_v01_stage-2_Best')
#     TRAIN_DATASET = Path('data/DETAILED_merged')
#     results_dir = 'data/MONS-Pelagic-Fish_results'
#     CRUISE_NAME = "MONS-Pelagic-Fish"
#     OSPAR = 'data/ospar_comp_au_2023_01_001-gis/ospar_comp_au_2023_01_001.shp' # From: https://odims.ospar.org/en/submissions/ospar_comp_au_2023_01/
#     DENSITY_CONSTANT = 340  # This constant is used in the R code for normalization into N per Liter (#/L)

#     # Step 5: Generate the Word document detailing the cruise
#     document_path = create_word_document(results_dir, OSPAR, CRUISE_NAME, DENSITY_CONSTANT, TRAIN_DATASET, MODEL_FILENAME)
#     print(f"Document generated at: {document_path}")

